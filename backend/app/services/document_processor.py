"""
Advanced Document Processor using docx2python
Provides superior DOCX processing compared to mammoth.js approach
"""

import re
import time
from typing import List, Dict, Any, Optional, Tuple
from pathlib import Path

# Core libraries
from docx2python import docx2python
from docx import Document
import xml.etree.ElementTree as ET

# Internal imports
from app.models.document import (
    DocumentAnalysis, DocumentType, Section, DocumentList, DocumentTable,
    Equation, DetectedElement, AuthorInfo, ProcessingConfig,
    ListType, SectionLevel, EquationType, ListItem, TableCell
)


class DocumentProcessor:
    """Advanced document processor leveraging docx2python for native DOCX access"""
    
    def __init__(self, config: ProcessingConfig):
        self.config = config
        self.paragraph_index = 0
        
    async def analyze_document(self, file_path: str, filename: str, file_size: int) -> DocumentAnalysis:
        """
        Analyze DOCX document using native Python processing
        
        Args:
            file_path: Path to the DOCX file
            filename: Original filename
            file_size: File size in bytes
            
        Returns:
            DocumentAnalysis: Complete document analysis
        """
        start_time = time.time()
        
        try:
            # Extract document content using docx2python
            doc_result = docx2python(file_path, extract_image=self.config.extract_images)
            
            # Also load with python-docx for style information
            doc = Document(file_path)
            
            # Extract all content
            paragraphs = self._extract_paragraphs(doc)
            word_styles = self._extract_styles(doc)
            
            # Analyze document structure
            title = self._detect_title(paragraphs, word_styles)
            authors = self._detect_authors(paragraphs, word_styles)
            abstract = self._detect_abstract(paragraphs, word_styles)
            keywords = self._detect_keywords(paragraphs, word_styles)
            
            # Extract structural elements
            sections = self._extract_sections(paragraphs, word_styles) if self.config.detect_equations else []
            lists = await self._extract_lists(doc, doc_result) if self.config.detect_lists else []
            tables = self._extract_tables(doc, doc_result) if self.config.detect_tables else []
            equations = await self._extract_equations(doc, doc_result) if self.config.detect_equations else []
            
            # Calculate metrics
            total_words = sum(len(p.text.split()) for p in paragraphs)
            confidence_score = self._calculate_confidence(
                title, authors, abstract, keywords, sections, lists, tables, equations
            )
            
            return DocumentAnalysis(
                filename=filename,
                file_size=file_size,
                document_type=DocumentType.DOCX,
                title=title,
                authors=authors,
                abstract=abstract,
                keywords=keywords,
                sections=sections,
                lists=lists,
                tables=tables,
                equations=equations,
                total_paragraphs=len(paragraphs),
                total_words=total_words,
                processing_time=time.time() - start_time,
                confidence_score=confidence_score,
                has_styles=len(word_styles) > 0,
                style_names=list(word_styles.keys())
            )
            
        except Exception as e:
            raise Exception(f"Failed to process document: {str(e)}")
    
    def _extract_paragraphs(self, doc: Document) -> List[Any]:
        """Extract all paragraphs from document"""
        return doc.paragraphs
    
    def _extract_styles(self, doc: Document) -> Dict[str, Any]:
        """Extract available styles from document"""
        styles = {}
        try:
            for style in doc.styles:
                if hasattr(style, 'name') and style.name:
                    styles[style.name] = style
        except:
            pass
        return styles
    
    def _detect_title(self, paragraphs: List[Any], styles: Dict[str, Any]) -> Optional[DetectedElement]:
        """Detect document title using Word styles and heuristics"""
        for i, para in enumerate(paragraphs[:10]):  # Check first 10 paragraphs
            if not para.text.strip():
                continue
                
            # Check for title styles
            style_name = para.style.name if para.style else None
            if style_name and 'title' in style_name.lower():
                return DetectedElement(
                    content=para.text.strip(),
                    confidence=0.95,
                    reasoning=f"Detected using Word style: {style_name}",
                    word_style=style_name,
                    paragraph_index=i
                )
            
            # Check for heading 1 as title if it's the first substantial content
            if style_name and 'heading 1' in style_name.lower() and i < 3:
                return DetectedElement(
                    content=para.text.strip(),
                    confidence=0.85,
                    reasoning="First Heading 1 style detected as title",
                    word_style=style_name,
                    paragraph_index=i
                )
        
        # Fallback: first non-empty paragraph
        for i, para in enumerate(paragraphs[:5]):
            if para.text.strip() and len(para.text.strip()) > 10:
                return DetectedElement(
                    content=para.text.strip(),
                    confidence=0.6,
                    reasoning="First substantial paragraph assumed as title",
                    paragraph_index=i
                )
        
        return None
    
    def _detect_authors(self, paragraphs: List[Any], styles: Dict[str, Any]) -> Optional[AuthorInfo]:
        """Detect author information using Word styles and patterns"""
        authors = AuthorInfo()
        
        for i, para in enumerate(paragraphs[:15]):  # Check first 15 paragraphs
            text = para.text.strip()
            if not text:
                continue
            
            style_name = para.style.name if para.style else None
            
            # Check for author-related styles
            if style_name and any(keyword in style_name.lower() for keyword in ['author', 'subtitle']):
                # Parse multiple authors
                if ',' in text or ' and ' in text:
                    names = re.split(r',|\band\b', text)
                    authors.names = [name.strip() for name in names if name.strip()]
                else:
                    authors.names = [text]
                break
            
            # Check for email patterns (likely author section)
            if '@' in text and '.' in text:
                # Extract emails
                email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
                emails = re.findall(email_pattern, text)
                authors.emails = emails
                
                # Also try to extract names from same or previous paragraphs
                if not authors.names and i > 0:
                    prev_text = paragraphs[i-1].text.strip()
                    if prev_text and not '@' in prev_text:
                        if ',' in prev_text or ' and ' in prev_text:
                            names = re.split(r',|\band\b', prev_text)
                            authors.names = [name.strip() for name in names if name.strip()]
                        else:
                            authors.names = [prev_text]
        
        return authors if authors.names or authors.emails else None
    
    def _detect_abstract(self, paragraphs: List[Any], styles: Dict[str, Any]) -> Optional[DetectedElement]:
        """Detect abstract using Word styles and keywords"""
        for i, para in enumerate(paragraphs):
            text = para.text.strip()
            if not text:
                continue
            
            style_name = para.style.name if para.style else None
            
            # Check for abstract styles
            if style_name and 'abstract' in style_name.lower():
                return DetectedElement(
                    content=text,
                    confidence=0.95,
                    reasoning=f"Detected using Word style: {style_name}",
                    word_style=style_name,
                    paragraph_index=i
                )
            
            # Check for abstract keyword
            if text.lower().startswith('abstract'):
                # Get the content (might be in same or next paragraph)
                content = text
                if text.lower().strip() == 'abstract' and i + 1 < len(paragraphs):
                    content = paragraphs[i + 1].text.strip()
                
                return DetectedElement(
                    content=content.replace('Abstract', '').replace('ABSTRACT', '').strip(),
                    confidence=0.9,
                    reasoning="Detected by 'Abstract' keyword",
                    paragraph_index=i
                )
        
        return None
    
    def _detect_keywords(self, paragraphs: List[Any], styles: Dict[str, Any]) -> Optional[DetectedElement]:
        """Detect keywords using patterns and styles"""
        keyword_patterns = [
            r'keywords?\s*[:\-]\s*(.+)',
            r'key\s*words?\s*[:\-]\s*(.+)',
            r'index\s*terms?\s*[:\-]\s*(.+)'
        ]
        
        for i, para in enumerate(paragraphs):
            text = para.text.strip()
            if not text:
                continue
            
            # Check each pattern
            for pattern in keyword_patterns:
                match = re.search(pattern, text, re.IGNORECASE)
                if match:
                    keywords = match.group(1).strip()
                    return DetectedElement(
                        content=keywords,
                        confidence=0.9,
                        reasoning=f"Detected using pattern: {pattern}",
                        paragraph_index=i
                    )
        
        return None
    
    def _extract_sections(self, paragraphs: List[Any], styles: Dict[str, Any]) -> List[Section]:
        """Extract document sections using Word heading styles"""
        sections = []
        section_id = 1
        
        for i, para in enumerate(paragraphs):
            text = para.text.strip()
            if not text:
                continue
            
            style_name = para.style.name if para.style else None
            if not style_name:
                continue
            
            # Check for heading styles
            level = None
            if 'heading 1' in style_name.lower():
                level = SectionLevel.SECTION
            elif 'heading 2' in style_name.lower():
                level = SectionLevel.SUBSECTION
            elif 'heading 3' in style_name.lower():
                level = SectionLevel.SUBSUBSECTION
            elif 'heading' in style_name.lower():
                level = SectionLevel.PARAGRAPH
            
            if level is not None:
                # Extract section number and title
                number_match = re.match(r'^(\d+(?:\.\d+)*)\s*\.?\s*(.+)', text)
                if number_match:
                    number = number_match.group(1)
                    title = number_match.group(2).strip()
                else:
                    number = str(section_id)
                    title = text
                
                # Find content until next section
                content = self._extract_section_content(paragraphs, i + 1, styles)
                
                sections.append(Section(
                    id=section_id,
                    number=number,
                    title=title,
                    content=content,
                    level=level,
                    confidence=0.9,
                    word_count=len(content.split()),
                    paragraph_start=i,
                    word_style=style_name
                ))
                
                section_id += 1
        
        return sections
    
    def _extract_section_content(self, paragraphs: List[Any], start_idx: int, styles: Dict[str, Any]) -> str:
        """Extract content between sections"""
        content_parts = []
        
        for i in range(start_idx, len(paragraphs)):
            para = paragraphs[i]
            style_name = para.style.name if para.style else None
            
            # Stop at next heading
            if style_name and 'heading' in style_name.lower():
                break
            
            text = para.text.strip()
            if text:
                content_parts.append(text)
        
        return '\n'.join(content_parts)
    
    async def _extract_lists(self, doc: Document, doc_result) -> List[DocumentList]:
        """Extract lists using Word's native list detection"""
        lists = []
        list_id = 1
        
        # Use python-docx to detect lists
        current_list = None
        list_items = []
        
        for i, para in enumerate(doc.paragraphs):
            # Check if paragraph is part of a list
            if self._is_list_paragraph(para):
                list_type = self._determine_list_type(para)
                
                if current_list is None or current_list != list_type:
                    # Save previous list if exists
                    if current_list is not None and list_items:
                        lists.append(self._create_document_list(
                            list_id, current_list, list_items
                        ))
                        list_id += 1
                    
                    # Start new list
                    current_list = list_type
                    list_items = []
                
                # Add item to current list
                item = self._create_list_item(para, list_type)
                if item:
                    list_items.append(item)
            else:
                # End current list
                if current_list is not None and list_items:
                    lists.append(self._create_document_list(
                        list_id, current_list, list_items
                    ))
                    list_id += 1
                    current_list = None
                    list_items = []
        
        # Don't forget the last list
        if current_list is not None and list_items:
            lists.append(self._create_document_list(
                list_id, current_list, list_items
            ))
        
        return lists
    
    def _is_list_paragraph(self, para) -> bool:
        """Check if paragraph is part of a list"""
        # Check for list numbering
        if hasattr(para, '_p') and para._p is not None:
            # Check for numbering properties in the paragraph
            num_pr = para._p.find('.//w:numPr', {
                'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
            })
            if num_pr is not None:
                return True
        
        # Fallback: check text patterns
        text = para.text.strip()
        if not text:
            return False
        
        # Check for bullet points or numbered items
        bullet_patterns = [
            r'^\s*[•·‣⁃]\s+',  # Bullet characters
            r'^\s*[-*+]\s+',    # Dash, asterisk, plus
            r'^\s*\d+\.\s+',    # Numbered items
            r'^\s*[a-zA-Z]\.\s+',  # Lettered items
            r'^\s*[ivxlcdm]+\.\s+',  # Roman numerals
        ]
        
        return any(re.match(pattern, text, re.IGNORECASE) for pattern in bullet_patterns)
    
    def _determine_list_type(self, para) -> ListType:
        """Determine the type of list"""
        text = para.text.strip()
        
        # Check for numbered patterns
        if re.match(r'^\s*\d+\.\s+', text):
            return ListType.ORDERED
        elif re.match(r'^\s*[a-zA-Z]\.\s+', text):
            return ListType.ORDERED
        elif re.match(r'^\s*[ivxlcdm]+\.\s+', text, re.IGNORECASE):
            return ListType.ORDERED
        else:
            return ListType.UNORDERED
    
    def _create_list_item(self, para, list_type: ListType) -> Optional[ListItem]:
        """Create a list item from paragraph"""
        text = para.text.strip()
        if not text:
            return None
        
        # Remove list markers
        clean_text = text
        level = 1
        index = None
        
        if list_type == ListType.ORDERED:
            # Extract number and content
            num_match = re.match(r'^\s*(\d+)\.\s+(.+)', text)
            if num_match:
                index = int(num_match.group(1))
                clean_text = num_match.group(2)
        else:
            # Remove bullet markers
            clean_text = re.sub(r'^\s*[•·‣⁃\-*+]\s+', '', text)
        
        # Determine nesting level (simplified)
        leading_spaces = len(text) - len(text.lstrip())
        if leading_spaces > 4:
            level = min(9, (leading_spaces // 4) + 1)
        
        return ListItem(
            content=clean_text,
            level=level,
            item_type=list_type.value,
            index=index,
            word_style=para.style.name if para.style else None
        )
    
    def _create_document_list(self, list_id: int, list_type: ListType, items: List[ListItem]) -> DocumentList:
        """Create a DocumentList from items"""
        max_depth = max(item.level for item in items) if items else 1
        is_nested = max_depth > 1
        
        return DocumentList(
            id=list_id,
            list_type=list_type,
            items=items,
            confidence=0.9,
            is_nested=is_nested,
            max_depth=max_depth
        )
    
    def _extract_tables(self, doc: Document, doc_result) -> List[DocumentTable]:
        """Extract tables using python-docx"""
        tables = []
        
        for i, table in enumerate(doc.tables):
            table_id = i + 1
            rows = len(table.rows)
            cols = len(table.columns) if table.rows else 0
            
            if rows == 0 or cols == 0:
                continue
            
            # Extract cell data
            cells = []
            for row_idx, row in enumerate(table.rows):
                cell_row = []
                for col_idx, cell in enumerate(row.cells):
                    cell_text = cell.text.strip()
                    is_header = row_idx == 0  # Simple heuristic
                    
                    table_cell = TableCell(
                        content=cell_text,
                        is_header=is_header
                    )
                    cell_row.append(table_cell)
                cells.append(cell_row)
            
            # Try to find table caption (look at paragraphs before/after)
            caption = self._find_table_caption(doc, i)
            
            tables.append(DocumentTable(
                id=table_id,
                rows=rows,
                columns=cols,
                cells=cells,
                caption=caption,
                confidence=0.95,
                has_headers=True,  # Assume first row is header
                paragraph_index=None  # TODO: Find actual position
            ))
        
        return tables
    
    def _find_table_caption(self, doc: Document, table_index: int) -> Optional[str]:
        """Find caption for a table by looking at nearby paragraphs"""
        # This is a simplified implementation
        # In practice, you'd need more sophisticated logic
        
        # Look for "Table X" patterns in paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if f"table {table_index + 1}" in text.lower():
                return text
        
        return None
    
    async def _extract_equations(self, doc: Document, doc_result) -> List[Equation]:
        """Extract mathematical equations from DOCX using OMML detection"""
        equations = []
        equation_id = 1
        
        # Extract OMML equations from document XML
        # This requires parsing the document's XML structure
        try:
            # Get document XML
            doc_xml = doc._document_part.blob
            root = ET.fromstring(doc_xml)
            
            # Find OMML equation elements
            namespaces = {
                'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
                'm': 'http://schemas.openxmlformats.org/officeDocument/2006/math'
            }
            
            # Find all math elements
            math_elements = root.findall('.//m:oMath', namespaces)
            
            for math_elem in math_elements:
                # Extract OMML content
                omml_content = ET.tostring(math_elem, encoding='unicode')
                
                # Convert OMML to text representation (simplified)
                text_content = self._omml_to_text(math_elem)
                
                # Try to convert to LaTeX (basic conversion)
                latex_equiv = self._omml_to_latex(math_elem)
                
                equations.append(Equation(
                    id=equation_id,
                    content=text_content,
                    equation_type=EquationType.OMML,
                    latex_equivalent=latex_equiv,
                    confidence=0.95,
                    omml_xml=omml_content,
                    is_display=True  # OMML equations are typically display equations
                ))
                
                equation_id += 1
                
        except Exception as e:
            # Fallback to text-based equation detection
            equations.extend(self._extract_text_equations(doc))
        
        return equations
    
    def _omml_to_text(self, math_elem) -> str:
        """Convert OMML element to readable text (simplified)"""
        # This is a basic implementation
        # In practice, you'd need more sophisticated OMML parsing
        return ''.join(math_elem.itertext()).strip()
    
    def _omml_to_latex(self, math_elem) -> Optional[str]:
        """Convert OMML to LaTeX (basic implementation)"""
        # This would require a proper OMML to LaTeX converter
        # For now, return a placeholder
        text = self._omml_to_text(math_elem)
        
        # Basic conversions
        latex = text
        latex = latex.replace('∑', '\\sum')
        latex = latex.replace('∫', '\\int')
        latex = latex.replace('α', '\\alpha')
        latex = latex.replace('β', '\\beta')
        latex = latex.replace('γ', '\\gamma')
        
        return latex if latex != text else None
    
    def _extract_text_equations(self, doc: Document) -> List[Equation]:
        """Fallback text-based equation extraction"""
        equations = []
        equation_id = 1000  # Different ID range for text equations
        
        for para in doc.paragraphs:
            text = para.text
            
            # Look for mathematical symbols and patterns
            math_patterns = [
                r'\$[^$]+\$',  # LaTeX inline
                r'\$\$[^$]+\$\$',  # LaTeX display
                r'[α-ωΑ-Ω]',  # Greek letters
                r'[∑∫∏√]',  # Math symbols
                r'\b\w+\s*=\s*\w+',  # Basic equations
            ]
            
            for pattern in math_patterns:
                matches = re.finditer(pattern, text)
                for match in matches:
                    content = match.group(0)
                    
                    equations.append(Equation(
                        id=equation_id,
                        content=content,
                        equation_type=EquationType.UNICODE_MATH,
                        confidence=0.7,
                        latex_equivalent=content.replace('$', '')
                    ))
                    
                    equation_id += 1
        
        return equations
    
    def _calculate_confidence(self, title, authors, abstract, keywords, 
                            sections, lists, tables, equations) -> float:
        """Calculate overall analysis confidence score"""
        scores = []
        
        # Element detection scores
        if title:
            scores.append(title.confidence)
        if authors:
            scores.append(0.8)  # AuthorInfo doesn't have confidence
        if abstract:
            scores.append(abstract.confidence)
        if keywords:
            scores.append(keywords.confidence)
        
        # Structural element scores
        if sections:
            scores.extend([s.confidence for s in sections])
        if lists:
            scores.extend([l.confidence for l in lists])
        if tables:
            scores.extend([t.confidence for t in tables])
        if equations:
            scores.extend([e.confidence for e in equations])
        
        return sum(scores) / len(scores) if scores else 0.5